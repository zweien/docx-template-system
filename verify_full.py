"""完整验证：模拟 render_report 流程，确认 PROMPT 被正确过滤"""
import sys, tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
sys.path.insert(0, "src")
from report_engine.prompt_parser import filter_prompt_paragraphs

# 1. 创建测试模板
with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    template_path = f.name

doc = Document()
doc.add_paragraph("{{PROJECT_NAME}}", style="Heading 2")
doc.add_paragraph("[[PROMPT: 立项依据: 请从国内外研究现状撰写 | mode=auto]]")
doc.add_paragraph("普通正文段落")
doc.add_paragraph("{%p if ENABLE_TEST %}")
doc.add_paragraph("{{p TEST_SUBDOC }}")
doc.add_paragraph("{%p endif %}")
doc.save(template_path)

print("=== 原始模板 ===")
for i, p in enumerate(Document(template_path).paragraphs):
    print(f"  [{i}] {p.text!r}")

# 2. 模拟 render_report 流程
tpl = DocxTemplate(template_path)
filter_prompt_paragraphs(tpl.get_docx())
print(f"\n过滤后 tpl.docx paragraphs: {len(tpl.docx.paragraphs)}")

# 模拟 render（只设置 is_rendered=True，不实际渲染）
tpl.is_rendered = True

# 3. 保存
with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    output_path = f.name
tpl.save(output_path)

# 4. 验证
doc_out = Document(output_path)
print(f"\n=== 输出文档（{len(doc_out.paragraphs)} 段）===")
for i, p in enumerate(doc_out.paragraphs):
    print(f"  [{i}] {p.text!r}")

all_text = "\n".join(p.text for p in doc_out.paragraphs)
assert "[[PROMPT:" not in all_text, "PROMPT 应被过滤"
assert "普通正文段落" in all_text, "普通段落应保留"
assert "{{PROJECT_NAME}}" in all_text, "占位符应保留（未渲染）"
print("\n✅ 验证通过：PROMPT 段落被正确过滤，且 save() 未重新加载")

# 5. 再验证：如果不设置 is_rendered=True，save() 会重新加载
tpl2 = DocxTemplate(template_path)
filter_prompt_paragraphs(tpl2.get_docx())
# 不设置 is_rendered=True
with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    output2 = f.name
tpl2.save(output2)
doc2 = Document(output2)
has_prompt = any("[[PROMPT:" in p.text for p in doc2.paragraphs)
print(f"\n不设置 is_rendered=True 时，save() 重新加载后仍含 PROMPT: {has_prompt}")
assert has_prompt, "预期会重新加载"
print("✅ 确认：save() 在 is_rendered=False 时会重新加载原始模板")

Path(template_path).unlink()
Path(output_path).unlink()
Path(output2).unlink()
