import sys, tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
sys.path.insert(0, "src")
from report_engine.prompt_parser import filter_prompt_paragraphs

with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    template_path = f.name

doc = Document()
doc.add_paragraph("{{PROJECT_NAME}}")
doc.add_paragraph("[[PROMPT: 测试: 提示内容]]")
doc.add_paragraph("普通段落")
doc.save(template_path)

tpl = DocxTemplate(template_path)
print(f"type(tpl) = {type(tpl)}")
print(f"tpl has element: {hasattr(tpl, 'element')}")
print(f"tpl paragraphs count: {len(tpl.paragraphs)}")
for p in tpl.paragraphs:
    print(f"  {p.text!r}")

# 直接用 tpl 本身（它是一个 Document-like 对象）
print("\n=== 直接传 tpl ===")
removed = filter_prompt_paragraphs(tpl)
print(f"删除数: {removed}")
print(f"过滤后 paragraphs 数: {len(tpl.paragraphs)}")
for p in tpl.paragraphs:
    print(f"  {p.text!r}")

with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    out = f.name
tpl.save(out)
doc_r = Document(out)
print(f"重新读取后 paragraphs 数: {len(doc_r.paragraphs)}")
for p in doc_r.paragraphs:
    print(f"  {p.text!r}")

Path(template_path).unlink()
Path(out).unlink()
