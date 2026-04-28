from __future__ import annotations

import tempfile
from pathlib import Path
from docx import Document

from report_engine.template_parser import parse_template


def _make_template_with_prompts(prompts_texts: list[str]) -> str:
    """Create a temporary docx template with given prompt paragraphs."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    doc = Document()
    doc.add_paragraph("{{PROJECT_NAME}}", style="Heading 2")
    for text in prompts_texts:
        doc.add_paragraph(text)
    doc.add_paragraph("{%p if ENABLE_RESEARCH_CONTENT %}")
    doc.add_paragraph("{{p RESEARCH_CONTENT_SUBDOC }}")
    doc.add_paragraph("{%p endif %}")
    doc.save(path)
    return path


def test_parse_template_extracts_prompts():
    path = _make_template_with_prompts([
        "[[PROMPT: 立项依据: 请从国内外研究现状撰写 | mode=auto]]",
        "[[PROMPT: 研究内容.技术路线: 描述技术路线 | mode=interactive]]",
    ])
    structure, warnings = parse_template(path)

    assert "prompts" in structure
    prompts = structure["prompts"]
    assert len(prompts) == 2

    assert prompts[0]["target"] == "立项依据"
    assert prompts[0]["prompt"] == "请从国内外研究现状撰写"
    assert prompts[0]["mode"] == "auto"
    assert prompts[0]["level"] == "section"

    assert prompts[1]["target"] == "研究内容.技术路线"
    assert prompts[1]["prompt"] == "描述技术路线"
    assert prompts[1]["mode"] == "interactive"
    assert prompts[1]["level"] == "paragraph"

    Path(path).unlink()


def test_parse_template_no_prompts():
    path = _make_template_with_prompts([])
    structure, warnings = parse_template(path)

    assert "prompts" in structure
    assert structure["prompts"] == []

    Path(path).unlink()


def test_parse_template_preserves_other_structure():
    path = _make_template_with_prompts([
        "[[PROMPT: 研究内容: 请撰写研究内容]]",
    ])
    structure, warnings = parse_template(path)

    assert "context_vars" in structure
    assert "sections" in structure
    assert len(structure["sections"]) == 1
    assert structure["sections"][0]["placeholder"] == "RESEARCH_CONTENT_SUBDOC"
    assert structure["prompts"][0]["target"] == "研究内容"

    Path(path).unlink()
