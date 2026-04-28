import sys, tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate

with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    template_path = f.name

doc = Document()
doc.add_paragraph("{{PROJECT_NAME}}")
doc.add_paragraph("[[PROMPT: 测试: 提示内容]]")
doc.add_paragraph("普通段落")
doc.save(template_path)

tpl = DocxTemplate(template_path)
print("DocxTemplate attributes:")
for attr in sorted(dir(tpl)):
    if not attr.startswith('_'):
        print(f"  {attr}")

print(f"\ntpl.docx = {tpl.docx}")
print(f"type(tpl) = {type(tpl)}")

# DocxTemplate inherits from DocxDocument which inherits from Document
print(f"isinstance(tpl, Document) = {isinstance(tpl, Document)}")
print(f"tpl.paragraphs = {tpl.paragraphs if hasattr(tpl, 'paragraphs') else 'N/A'}")

# try get_docx
print(f"tpl.get_docx() = {tpl.get_docx()}")
print(f"tpl.get_docx() is tpl = {tpl.get_docx() is tpl}")

# check element
print(f"tpl.element = {tpl.element}")
print(f"tpl.element.body = {tpl.element.body}")

Path(template_path).unlink()
