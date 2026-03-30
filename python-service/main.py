"""Lightweight Python document generation service.

Replaces {{ key }} placeholders in .docx templates with form data values.
"""

import re
import os
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel

CJK = r"\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff"

block_start_pattern = re.compile(r"\{\{\s*#([\w" + CJK + r"]+)\s*\}\}")
block_end_pattern = re.compile(r"\{\{\s*/([\w" + CJK + r"]+)\s*\}\}")
placeholder_pattern = re.compile(r"\{\{\s*([\w" + CJK + r"]+)\s*\}\}")
choice_control_pattern = re.compile(r"\{\{\s*选项:([\w" + CJK + r"]+)\|(single|multiple)\s*\}\}")
choice_option_pattern = re.compile(r"^([□☐☑])\s*(.+)$")
checked_sym_chars = {"0052"}
unchecked_sym_chars = {"00A3"}

app = FastAPI(title="DOCX Template Generator", version="1.0.0")


class GenerateRequest(BaseModel):
    template_path: str
    output_filename: str
    form_data: dict[str, Any] = {}


@app.get("/health")
async def health():
    return {"status": "ok"}


def replace_placeholders_in_paragraph(paragraph, form_data):
    """Replace {{ key }} placeholders in a paragraph's full text."""
    full_text = paragraph.text
    if "{{" not in full_text:
        return

    # Build replacement map
    def replacer(match):
        key = match.group(1)
        value = form_data.get(key)
        if value is None:
            return match.group(0)
        return str(value)

    new_text = placeholder_pattern.sub(replacer, full_text)

    if new_text == full_text:
        return

    # Replace text while trying to preserve formatting of the first run
    if paragraph.runs:
        # Clear all runs except the first
        first_run = paragraph.runs[0]
        first_run.text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


def replace_paragraph_text(paragraph, new_text: str):
    """Replace paragraph text while keeping the first run container."""
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


def replace_placeholders_in_table(table, form_data):
    """Replace placeholders in table cells."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, form_data)


def process_choice_blocks(doc, form_data):
    """Process {{选项:key|single}} / {{选项:key|multiple}} blocks in paragraphs."""
    paragraphs = doc.paragraphs
    index = 0

    while index < len(paragraphs):
        control_match = choice_control_pattern.search(paragraphs[index].text)
        if not control_match:
            index += 1
            continue

        key = control_match.group(1)
        mode = control_match.group(2)
        raw_value = form_data.get(key)

        if mode == "multiple" and isinstance(raw_value, list):
            selected_values = {str(item) for item in raw_value}
        elif raw_value not in (None, ""):
            selected_values = {str(raw_value)}
        else:
            selected_values = set()

        replace_paragraph_text(paragraphs[index], "")
        index += 1

        while index < len(paragraphs):
            option_match = choice_option_pattern.match(paragraphs[index].text)
            if not option_match:
                break

            label = option_match.group(2)
            marker = "☑" if label in selected_values else "☐"
            replace_paragraph_text(paragraphs[index], f"{marker} {label}")
            index += 1


def extract_inline_choice_segments(paragraph):
    segments: list[dict[str, str]] = []
    current_marker = None
    text_buffer = ""

    for run in paragraph.runs:
        run_text = run.text or ""
        sym_chars = [
            child.get(qn("w:char")).upper()
            for child in run._r
            if child.tag == qn("w:sym") and child.get(qn("w:char"))
        ]

        if sym_chars:
            if current_marker and text_buffer.strip():
                segments.append({"label": text_buffer.strip(), "marker": current_marker})
                text_buffer = ""

            for sym_char in sym_chars:
                if sym_char in checked_sym_chars:
                    current_marker = "0052"
                elif sym_char in unchecked_sym_chars:
                    current_marker = "00A3"
            continue

        if current_marker:
            text_buffer += run_text

    if current_marker and text_buffer.strip():
        segments.append({"label": text_buffer.strip(), "marker": current_marker})

    return segments


def update_paragraph_sym_markers(paragraph, next_markers: list[str]) -> None:
    marker_index = 0
    for run in paragraph.runs:
        for child in run._r:
            if child.tag != qn("w:sym"):
                continue

            current_char = child.get(qn("w:char"))
            if not current_char:
                continue

            current_char = current_char.upper()
            if current_char not in checked_sym_chars and current_char not in unchecked_sym_chars:
                continue

            if marker_index >= len(next_markers):
                return

            child.set(qn("w:char"), next_markers[marker_index])
            marker_index += 1


def process_inline_choice_paragraphs(doc, form_data):
    """Process inline paragraphs like '单项：☑是☐否' stored as w:sym runs."""
    for paragraph in doc.paragraphs:
        text = paragraph.text
        separator_index = text.find("：") if "：" in text else text.find(":")
        if separator_index < 0:
            continue

        key = text[:separator_index].strip()
        if not key or key not in form_data:
            continue

        segments = extract_inline_choice_segments(paragraph)
        if len(segments) < 2:
            continue

        raw_value = form_data.get(key)
        if isinstance(raw_value, list):
            selected_values = {str(item) for item in raw_value}
        elif raw_value not in (None, ""):
            selected_values = {str(raw_value)}
        else:
            selected_values = set()

        next_markers = [
            "0052" if segment["label"] in selected_values else "00A3"
            for segment in segments
        ]
        update_paragraph_sym_markers(paragraph, next_markers)


def process_table_block(table, block_name, rows_data, form_data):
    """Process {{#name}}...{{/name}} blocks in a table."""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    table_element = table._tbl

    # Find block marker rows
    start_row_idx = None
    end_row_idx = None
    for i, row in enumerate(table.rows):
        row_text = "".join(cell.text for cell in row.cells)
        if block_start_pattern.search(row_text):
            start_row_idx = i
        if block_end_pattern.search(row_text):
            end_row_idx = i
            break

    if start_row_idx is None or end_row_idx is None:
        return

    # Extract template rows (between markers, exclusive)
    tr_elements = table_element.findall(f".//{ns}tr")
    template_rows_xml = [deepcopy(tr_elements[i]) for i in range(start_row_idx + 1, end_row_idx)]

    # Remove marker rows and template rows (reverse order to preserve indices)
    for idx in reversed(range(start_row_idx, end_row_idx + 1)):
        parent = tr_elements[idx].getparent()
        if parent is not None:
            parent.remove(tr_elements[idx])

    if not rows_data:
        return

    # Find anchor for insertion
    tr_elements = table_element.findall(f".//{ns}tr")
    anchor_idx = min(start_row_idx, len(tr_elements)) if tr_elements else 0
    anchor = tr_elements[anchor_idx] if anchor_idx < len(tr_elements) else None
    insertion_point = None

    # Clone and insert rows for each data item
    for item in rows_data:
        for row_xml in template_rows_xml:
            new_row = deepcopy(row_xml)
            _replace_in_row_xml(new_row, item)
            if insertion_point is not None:
                insertion_point.addnext(new_row)
                insertion_point = new_row
            elif anchor is not None:
                anchor.addprevious(new_row)
                insertion_point = new_row
            else:
                table_element.append(new_row)
                insertion_point = new_row


def _replace_in_row_xml(row_xml, data):
    """Replace placeholders in a cloned row XML element."""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for paragraph in row_xml.findall(f".//{ns}p"):
        full_text = "".join((t.text or "") for t in paragraph.findall(f".//{ns}t"))
        if "{{" not in full_text:
            continue

        new_text = placeholder_pattern.sub(
            lambda m: str(data.get(m.group(1), m.group(0))),
            full_text
        )
        if new_text == full_text:
            continue

        runs = paragraph.findall(f"{ns}r")
        if runs:
            t_elem = runs[0].find(f"{ns}t")
            if t_elem is not None:
                t_elem.text = new_text
            for run in runs[1:]:
                for t in run.findall(f"{ns}t"):
                    t.text = ""


@app.post("/generate")
async def generate_document(request: GenerateRequest):
    template_path = Path(request.template_path)

    if not template_path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"Template file not found: {template_path}",
        )

    if not template_path.suffix.lower() == ".docx":
        raise HTTPException(
            status_code=400,
            detail="Only .docx files are supported",
        )

    try:
        doc = Document(str(template_path))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to open document: {e}",
        )

    simple_data: dict[str, str] = {}
    choice_data: dict[str, str | list[str]] = {}
    table_data: dict[str, list[dict[str, str]]] = {}
    for key, value in request.form_data.items():
        if isinstance(value, list) and value and all(isinstance(item, dict) for item in value):
            table_data[key] = value
        elif isinstance(value, list):
            choice_data[key] = [str(item) for item in value]
        else:
            simple_data[key] = str(value)

    # Replace placeholders in body paragraphs
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, simple_data)

    process_choice_blocks(doc, choice_data | simple_data)
    process_inline_choice_paragraphs(doc, choice_data | simple_data)

    # Collect all block names present in any table
    all_block_names: set[str] = set(table_data.keys())
    for table in doc.tables:
        for row in table.rows:
            row_text = "".join(cell.text for cell in row.cells)
            for match in block_start_pattern.finditer(row_text):
                all_block_names.add(match.group(1))

    # Replace placeholders in tables and process blocks
    for table in doc.tables:
        replace_placeholders_in_table(table, simple_data)
        for block_name in all_block_names:
            process_table_block(table, block_name, table_data.get(block_name, []), simple_data)

    # Save to a temp file and return
    output_dir = Path(os.environ.get("OUTPUT_DIR", "/tmp/docx-output"))
    output_dir.mkdir(parents=True, exist_ok=True)

    output_path = output_dir / request.output_filename
    doc.save(str(output_path))

    return FileResponse(
        str(output_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=request.output_filename,
    )


if __name__ == "__main__":
    import uvicorn

    port = int(os.environ.get("PORT", "8065"))
    uvicorn.run(app, host="0.0.0.0", port=port)
