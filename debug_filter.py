import sys, tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
sys.path.insert(0, "src")
from report_engine.prompt_parser import filter_prompt_paragraphs, PROMPT_PREFIX, PROMPT_SUFFIX

with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    template_path = f.name

doc = Document()
doc.add_paragraph("{{PROJECT_NAME}}")
doc.add_paragraph("[[PROMPT: 测试: 提示内容]]")
doc.add_paragraph("普通段落")
doc.save(template_path)

# 方法1: 直接 Document + filter
doc1 = Document(template_path)
print("=== 方法1: 直接 Document ===")
print(f"过滤前 paragraphs 数: {len(doc1.paragraphs)}")
for p in doc1.paragraphs:
    print(f"  text={p.text!r}, elem={p._element.tag}")

removed1 = filter_prompt_paragraphs(doc1)
print(f"删除数: {removed1}")
print(f"过滤后 paragraphs 数: {len(doc1.paragraphs)}")
for p in doc1.paragraphs:
    print(f"  text={p.text!r}")

# 保存并重新读取
with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    out1 = f.name
doc1.save(out1)
doc1_r = Document(out1)
print(f"重新读取后 paragraphs 数: {len(doc1_r.paragraphs)}")
for p in doc1_r.paragraphs:
    print(f"  text={p.text!r}")

print()

# 方法2: DocxTemplate + filter
tpl = DocxTemplate(template_path)
print("=== 方法2: DocxTemplate ===")
print(f"过滤前 paragraphs 数: {len(tpl.paragraphs)}")
removed2 = filter_prompt_paragraphs(tpl.get_docx())
print(f"删除数: {removed2}")
print(f"过滤后 paragraphs 数: {len(tpl.paragraphs)}")
for p in tpl.paragraphs:
    print(f"  text={p.text!r}")

# 保存并重新读取
with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    out2 = f.name
tpl.save(out2)
doc2_r = Document(out2)
print(f"重新读取后 paragraphs 数: {len(doc2_r.paragraphs)}")
for p in doc2_r.paragraphs:
    print(f"  text={p.text!r}")

# 检查 XML
print("\n=== 检查 body 子元素 ===")
body = doc2_r.element.body
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    text = "".join(t.text or "" for t in child.findall(".//w:t", namespaces=doc2_r.element.nsmap))
    print(f"  {tag}: {text[:50]!r}")

Path(template_path).unlink()
Path(out1).unlink()
Path(out2).unlink()
