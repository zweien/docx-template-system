import sys, tempfile
from pathlib import Path
from docx import Document
from docxtpl import DocxTemplate
sys.path.insert(0, "src")
from report_engine.prompt_parser import filter_prompt_paragraphs

with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    template_path = f.name

doc = Document()
doc.add_paragraph("{{PROJECT_NAME}}")
doc.add_paragraph("[[PROMPT: 测试: 提示内容]]")
doc.add_paragraph("普通段落")
doc.save(template_path)

tpl = DocxTemplate(template_path)
print(f"tpl.docx before get_docx = {tpl.docx}")

docx = tpl.get_docx()
print(f"tpl.docx after get_docx = {tpl.docx}")
print(f"docx is tpl.docx = {docx is tpl.docx}")
print(f"docx paragraphs count: {len(docx.paragraphs)}")
for p in docx.paragraphs:
    print(f"  {p.text!r}")

removed = filter_prompt_paragraphs(docx)
print(f"删除数: {removed}")
print(f"过滤后 paragraphs 数: {len(docx.paragraphs)}")
for p in docx.paragraphs:
    print(f"  {p.text!r}")

with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
    out = f.name
tpl.save(out)
doc_r = Document(out)
print(f"重新读取后 paragraphs 数: {len(doc_r.paragraphs)}")
for p in doc_r.paragraphs:
    print(f"  {p.text!r}")

Path(template_path).unlink()
Path(out).unlink()
