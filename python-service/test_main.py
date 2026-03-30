import unittest
import sys
import types

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

fastapi_stub = types.ModuleType("fastapi")


class _FastAPI:
    def __init__(self, *args, **kwargs):
        pass

    def get(self, *args, **kwargs):
        def decorator(func):
            return func
        return decorator

    def post(self, *args, **kwargs):
        def decorator(func):
            return func
        return decorator


class _HTTPException(Exception):
    pass


fastapi_stub.FastAPI = _FastAPI
fastapi_stub.HTTPException = _HTTPException
sys.modules.setdefault("fastapi", fastapi_stub)

responses_stub = types.ModuleType("fastapi.responses")
responses_stub.FileResponse = object
sys.modules.setdefault("fastapi.responses", responses_stub)

pydantic_stub = types.ModuleType("pydantic")


class _BaseModel:
    pass


pydantic_stub.BaseModel = _BaseModel
sys.modules.setdefault("pydantic", pydantic_stub)

from main import (
    process_choice_blocks,
    process_choice_blocks_in_paragraphs,
    process_inline_choice_paragraphs,
    process_table_block,
)


def set_vertical_merge(cell, value: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    v_merge = OxmlElement("w:vMerge")
    v_merge.set(qn("w:val"), value)
    tc_pr.append(v_merge)


def get_vertical_merge_from_row(row, cell_index: int = 0) -> str | None:
    tc_pr = row._tr.tc_lst[cell_index].tcPr
    if tc_pr is None or tc_pr.vMerge is None:
        return None
    return tc_pr.vMerge.val or "continue-empty"


def append_sym_run(paragraph, char: str) -> None:
    run = paragraph.add_run()
    sym = OxmlElement("w:sym")
    sym.set(qn("w:font"), "Wingdings 2")
    sym.set(qn("w:char"), char)
    run._r.append(sym)


def get_sym_chars(paragraph) -> list[str]:
    chars: list[str] = []
    for run in paragraph.runs:
        for child in run._r:
            if child.tag == qn("w:sym"):
                chars.append(child.get(qn("w:char")))
    return chars


class ProcessTableBlockTest(unittest.TestCase):
    def test_inserts_rows_before_following_section_and_preserves_merge_chain(self) -> None:
        doc = Document()
        table = doc.add_table(rows=5, cols=2)

        table.cell(0, 0).text = "左侧合并标题"
        table.cell(0, 1).text = "表头"
        set_vertical_merge(table.cell(0, 0), "restart")

        table.cell(1, 0).text = ""
        table.cell(1, 1).text = "{{#工作计划及完成情况}}"
        set_vertical_merge(table.cell(1, 0), "continue")

        table.cell(2, 0).text = ""
        table.cell(2, 1).text = "{{计划任务}}"
        set_vertical_merge(table.cell(2, 0), "continue")

        table.cell(3, 0).text = ""
        table.cell(3, 1).text = "{{/工作计划及完成情况}}"
        set_vertical_merge(table.cell(3, 0), "continue")

        table.cell(4, 0).text = "下一节"
        table.cell(4, 1).text = "后续内容"

        process_table_block(
            table,
            "工作计划及完成情况",
            [{"计划任务": "第一行"}, {"计划任务": "第二行"}],
            {},
        )

        self.assertEqual(len(table.rows), 4)
        self.assertEqual(table.rows[0].cells[1].text, "表头")
        self.assertEqual(table.rows[1].cells[1].text, "第一行")
        self.assertEqual(table.rows[2].cells[1].text, "第二行")
        self.assertEqual(table.rows[3].cells[0].text, "下一节")

        self.assertEqual(get_vertical_merge_from_row(table.rows[0]), "restart")
        self.assertEqual(get_vertical_merge_from_row(table.rows[1]), "continue")
        self.assertEqual(get_vertical_merge_from_row(table.rows[2]), "continue")
        self.assertIsNone(get_vertical_merge_from_row(table.rows[3]))


class ProcessChoiceBlocksTest(unittest.TestCase):
    def test_single_choice_replaces_only_marker_and_clears_control_line(self) -> None:
        doc = Document()
        doc.add_paragraph("{{选项:性别|single}}")
        doc.add_paragraph("□ 男")
        doc.add_paragraph("□ 女")

        process_choice_blocks(doc, {"性别": "女"})

        self.assertEqual(len(doc.paragraphs), 2)
        self.assertEqual(doc.paragraphs[0].text, "☐ 男")
        self.assertEqual(doc.paragraphs[1].text, "☑ 女")

    def test_multiple_choice_replaces_only_selected_markers(self) -> None:
        doc = Document()
        doc.add_paragraph("{{选项:爱好|multiple}}")
        doc.add_paragraph("□ 篮球")
        doc.add_paragraph("□ 音乐")

        process_choice_blocks(doc, {"爱好": ["音乐"]})

        self.assertEqual(len(doc.paragraphs), 2)
        self.assertEqual(doc.paragraphs[0].text, "☐ 篮球")
        self.assertEqual(doc.paragraphs[1].text, "☑ 音乐")

    def test_control_line_can_drive_inline_wingdings_choices(self) -> None:
        doc = Document()
        doc.add_paragraph("{{选项:单项|single}}")

        inline = doc.add_paragraph()
        inline.add_run("单项：")
        append_sym_run(inline, "0052")
        inline.add_run("是")
        append_sym_run(inline, "00A3")
        inline.add_run("否")

        process_choice_blocks(doc, {"单项": "否"})

        self.assertEqual(len(doc.paragraphs), 1)
        self.assertEqual(get_sym_chars(doc.paragraphs[0]), ["00A3", "0052"])

    def test_table_cell_control_line_can_drive_inline_wingdings_choices(self) -> None:
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        cell.paragraphs[0].text = "{{选项:五类|single}}"
        inline = cell.add_paragraph()
        inline.add_run("符合“五类”来源 ")
        append_sym_run(inline, "00A3")
        inline.add_run("是   ")
        append_sym_run(inline, "00A3")
        inline.add_run("否")

        process_choice_blocks_in_paragraphs(cell.paragraphs, {"五类": "是"})

        self.assertEqual(len(cell.paragraphs), 1)
        self.assertEqual(get_sym_chars(cell.paragraphs[0]), ["0052", "00A3"])

    def test_inline_choice_paragraphs_should_update_wingdings_symbols(self) -> None:
        doc = Document()

        single = doc.add_paragraph()
        single.add_run("单项：")
        append_sym_run(single, "0052")
        single.add_run("是")
        append_sym_run(single, "00A3")
        single.add_run("否")

        multiple = doc.add_paragraph()
        multiple.add_run("多选：")
        append_sym_run(multiple, "00A3")
        multiple.add_run("选项1")
        append_sym_run(multiple, "0052")
        multiple.add_run("选项2")
        append_sym_run(multiple, "0052")
        multiple.add_run("选项3")
        append_sym_run(multiple, "00A3")
        multiple.add_run("选项4")

        process_inline_choice_paragraphs(doc, {"单项": "否", "多选": ["选项1", "选项4"]})

        self.assertEqual(get_sym_chars(doc.paragraphs[0]), ["00A3", "0052"])
        self.assertEqual(get_sym_chars(doc.paragraphs[1]), ["0052", "00A3", "00A3", "0052"])


if __name__ == "__main__":
    unittest.main()
