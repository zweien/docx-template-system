#!/usr/bin/env python3
"""从简化内容描述构建合法的 report-engine payload。

用法:
    # 标准模式：从 content.json 构建 payload
    python build_payload.py --content content.json --output payload.json --template template.docx

    # 仅列出模板中的 prompts
    python build_payload.py --template template.docx --list-prompts

    # 自动根据模板 prompts 生成内容（需配置 AI_BASE_URL / AI_API_KEY / AI_MODEL）
    python build_payload.py --template template.docx --output payload.json --auto-generate \
        --context '{"PROJECT_NAME": "XXX项目", "APPLICANT_ORG": "XX大学"}'

content.json 格式（比完整 payload 简单得多）:
{
  "title": "报告标题",
  "org": "单位名称",
  "leader": "负责人",
  "period": "2026年1月-12月",
  "sections": [
    {
      "name": "研究内容与技术路线",
      "blocks": [
        {"type": "paragraph", "text": "本项目旨在..."},
        {"type": "table", "headers": ["指标","目标"], "rows": [["准确率","95%"]]}
      ]
    }
  ],
  "attachments": [
    {
      "name": "经费预算",
      "blocks": [
        {"type": "appendix_table", "headers": ["科目","金额"], "rows": [["设备费","50万"]]}
      ]
    }
  ]
}

脚本自动完成:
  - 从 name 生成 id/placeholder/flag_name
  - 合并默认 style_map
  - 添加 attachments_bundle 配置
  - 校验 block 字段完整性
  - 输出合法的 payload JSON
"""

import argparse
import json
import re
import sys
from pathlib import Path


# ── 默认 style_map ─────────────────────────────────────────

DEFAULT_STYLE_MAP = {
    "table": "ResearchTable",
    "appendix_table": "AppendixTable",
    "figure_paragraph": "Figure Paragraph",
    "legend": "Legend",
    "note": "Note",
    "quote": "Quote",
    "checklist": "Checklist",
    "code_block": "CodeBlock",
}

# ── block 必填字段 ─────────────────────────────────────────

BLOCK_REQUIRED = {
    "heading": ["text"],
    "paragraph": ["text"],
    "bullet_list": ["items"],
    "numbered_list": ["items"],
    "table": ["headers", "rows"],
    "image": ["path"],
    "page_break": [],
    "rich_paragraph": ["segments"],
    "note": ["text"],
    "quote": ["text"],
    "two_images_row": ["images"],
    "appendix_table": ["headers", "rows"],
    "checklist": ["items"],
    "horizontal_rule": [],
    "toc_placeholder": [],
    "code_block": ["code"],
    "formula": ["latex"],
    "columns": ["count", "columns"],
    "mermaid": ["code"],
}


# 常见中文章节名到标准 id 的映射
SECTION_NAME_MAP = {
    "项目概述": "project_overview",
    "研究内容": "research_content",
    "研究内容与技术路线": "research_content",
    "技术路线": "technical_route",
    "研究基础": "research_basis",
    "研究基础与条件保障": "research_basis",
    "实施计划": "implementation_plan",
    "实施计划与进度安排": "implementation_plan",
    "经费预算": "budget",
    "参考文献": "references",
    "文献综述": "literature_review",
    "研究方法": "methodology",
    "预期成果": "expected_results",
    "目录": "toc",
}


def name_to_id(name: str) -> str:
    """将中文名称转为 snake_case id。"""
    # 去掉中文数字前缀
    cleaned = re.sub(r"^[一二三四五六七八九十]+[、．.]\s*", "", name).strip()

    # 先查映射表
    for key, val in SECTION_NAME_MAP.items():
        if key in cleaned:
            return val

    # 如果是纯中文，返回 None 由调用方决定
    if re.search(r"[\u4e00-\u9fff]", cleaned):
        return None

    # 英文：转 snake_case
    return re.sub(r"[^a-z0-9]+", "_", cleaned.lower()).strip("_")


def build_section(section: dict, index: int) -> dict:
    """构建一个 section 对象。"""
    name = section.get("name", f"Section {index + 1}")
    blocks = section.get("blocks", [])

    # 允许用户手动指定 placeholder/flag_name/id
    sid = section.get("id") or name_to_id(name)
    if sid is None:
        sid = f"section_{index + 1}"

    placeholder = section.get("placeholder") or f"{sid.upper()}_SUBDOC"
    flag_name = section.get("flag_name") or f"ENABLE_{sid.upper()}"

    # 校验 blocks
    for i, block in enumerate(blocks):
        btype = block.get("type")
        if not btype:
            raise ValueError(f"Section '{name}' block[{i}]: missing 'type'")
        if btype not in BLOCK_REQUIRED:
            raise ValueError(f"Section '{name}' block[{i}]: unknown type '{btype}'")
        missing = [f for f in BLOCK_REQUIRED[btype] if f not in block]
        if missing:
            raise ValueError(
                f"Section '{name}' block[{i}] ({btype}): missing fields: {', '.join(missing)}"
            )

    return {
        "id": sid,
        "placeholder": placeholder,
        "flag_name": flag_name,
        "enabled": section.get("enabled", True),
        "blocks": blocks,
        "order": index + 1,
    }


def build_attachment(attachment: dict, index: int) -> dict:
    """构建一个 attachment 对象。"""
    name = attachment.get("name", f"Appendix {index + 1}")
    blocks = attachment.get("blocks", [])

    aid = f"appendix_{index + 1}"
    placeholder = f"APPENDIX_{index + 1}_SUBDOC"
    flag_name = f"ENABLE_APPENDIX_{index + 1}"

    # 校验 blocks
    for i, block in enumerate(blocks):
        btype = block.get("type")
        if not btype:
            raise ValueError(f"Attachment '{name}' block[{i}]: missing 'type'")
        if btype not in BLOCK_REQUIRED:
            raise ValueError(f"Attachment '{name}' block[{i}]: unknown type '{btype}'")
        missing = [f for f in BLOCK_REQUIRED[btype] if f not in block]
        if missing:
            raise ValueError(
                f"Attachment '{name}' block[{i}] ({btype}): missing fields: {', '.join(missing)}"
            )

    return {
        "id": aid,
        "placeholder": placeholder,
        "flag_name": flag_name,
        "enabled": True,
        "title": name,
        "title_level": 2,
        "blocks": blocks,
        "order": index + 1,
    }


def _resolve_style_map(template_path: str) -> dict:
    """根据模板中实际存在的样式生成适配的 style_map。"""
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document(template_path)
        paragraph_styles = set()
        table_styles = set()
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                paragraph_styles.add(style.name)
            elif style.type == WD_STYLE_TYPE.TABLE:
                table_styles.add(style.name)
    except Exception:
        return dict(DEFAULT_STYLE_MAP)

    # 样式候选列表（按优先级）
    STYLE_FALLBACKS = {
        "heading_1": ["Heading 1"],
        "heading_2": ["Heading 2", "Heading 1"],
        "heading_3": ["Heading 3", "Heading 2", "Heading 1"],
        "heading_4": ["Heading 4", "Heading 3", "Heading 2", "Heading 1"],
        "heading_5": ["Heading 5", "Heading 4", "Heading 3", "Heading 2", "Heading 1"],
        "body": ["Body Text", "Normal"],
        "table_caption": ["TableCaption", "Caption", "Body Text", "Normal"],
        "figure_caption": ["FigureCaption", "Caption", "Body Text", "Normal"],
        "legend": ["Legend", "Body Text", "Normal"],
        "figure_paragraph": ["Figure Paragraph", "Body Text", "Normal"],
        "table": ["ResearchTable", "Table Grid", "Normal Table", "Table Normal"],
        "bullet_list": ["List Bullet", "List Bullet 2", "List Paragraph", "Body Text", "Normal"],
        "numbered_list": ["List Number", "List Number 2", "List Paragraph", "Body Text", "Normal"],
        "note": ["Note", "Body Text", "Normal"],
        "quote": ["Quote", "Intense Quote", "Body Text", "Normal"],
        "checklist": ["Checklist", "List Bullet", "List Paragraph", "Body Text", "Normal"],
        "code_block": ["CodeBlock", "Body Text", "Normal"],
        "appendix_table": ["AppendixTable", "Table Grid", "Normal Table", "Table Normal"],
    }

    style_map = {}
    for key, candidates in STYLE_FALLBACKS.items():
        is_table = key in ("table", "appendix_table")
        available = table_styles if is_table else paragraph_styles
        for cand in candidates:
            if cand in available:
                style_map[key] = cand
                break

    # 合并用户自定义（如果有）
    return style_map


def build_payload(content: dict, template_path: str = None) -> dict:
    """从简化内容描述构建完整 payload。"""
    # context
    context = {}
    if content.get("title"):
        context["PROJECT_NAME"] = content["title"]
    if content.get("org"):
        context["APPLICANT_ORG"] = content["org"]
    if content.get("leader"):
        context["PROJECT_LEADER"] = content["leader"]
    if content.get("period"):
        context["PROJECT_PERIOD"] = content["period"]
    # 合并用户提供的额外 context
    if "extra_context" in content:
        context.update(content["extra_context"])

    # sections
    sections = []
    for i, sec in enumerate(content.get("sections", [])):
        sections.append(build_section(sec, i))

    # attachments
    attachments = []
    for i, att in enumerate(content.get("attachments", [])):
        attachments.append(build_attachment(att, i))

    # attachments_bundle
    has_attachments = len(attachments) > 0
    attachments_bundle = {
        "enabled": has_attachments,
        "placeholder": "APPENDICES_SUBDOC",
        "flag_name": "ENABLE_APPENDICES",
        "page_break_between_attachments": True,
        "include_attachment_title": True,
    }

    # style_map
    if template_path:
        style_map = _resolve_style_map(template_path)
    else:
        style_map = dict(DEFAULT_STYLE_MAP)
    if content.get("style_map"):
        style_map.update(content["style_map"])

    return {
        "context": context,
        "sections": sections,
        "attachments": attachments,
        "attachments_bundle": attachments_bundle,
        "style_map": style_map,
    }


# ── 模板 prompt 相关功能 ─────────────────────────────────────────

def extract_prompts_from_template(template_path: str) -> list:
    """从模板中提取 PROMPT 注释。"""
    # 复用 analyze_template.py 的逻辑
    script_dir = Path(__file__).resolve().parent
    sys.path.insert(0, str(script_dir))
    try:
        from analyze_template import analyze_template
        result = analyze_template(template_path)
        return result.get("prompts", [])
    except ImportError:
        # 备用：直接解析
        from docx import Document
        import re as _re
        doc = Document(template_path)
        prompts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            match = _re.match(r"\[\[PROMPT:\s*([^:]+):\s*(.+?)\s*(?:\|\s*mode=(\w+))?\s*\]\]", text)
            if match:
                prompts.append({
                    "target": match.group(1).strip(),
                    "prompt": match.group(2).strip(),
                    "mode": match.group(3) if match.group(3) else "auto",
                    "level": "paragraph" if "." in match.group(1) else "section",
                })
        return prompts


def list_prompts(template_path: str) -> None:
    """列出模板中的所有 prompts。"""
    prompts = extract_prompts_from_template(template_path)
    if not prompts:
        print(f"模板 {template_path} 中没有发现 PROMPT 注释")
        return

    print(f"模板: {template_path}")
    print(f"发现 {len(prompts)} 个 PROMPT 注释:\n")
    for i, p in enumerate(prompts, 1):
        level_icon = "S" if p["level"] == "section" else "P"
        mode_icon = "A" if p["mode"] == "auto" else "I"
        print(f"  [{i}] [{level_icon}/{mode_icon}] {p['target']}")
        print(f"      {p['prompt']}")
        print()


# 扩展映射：英文 section id -> 中文标题
SECTION_ID_TO_CHINESE = {
    "project_overview": "项目概述",
    "research_content": "研究内容",
    "technical_route": "技术路线",
    "implementation_plan": "实施计划",
    "expected_results": "预期成果",
    "research_basis": "研究基础",
    "budget": "经费预算",
    "references": "参考文献",
    "literature_review": "文献综述",
    "methodology": "研究方法",
    "toc": "目录",
}


def _build_sections_from_flags(flags: list, subdocs: list) -> list:
    """从 flags 和 subdocs 配对构建 sections（复用 template_parser 逻辑）。"""
    sections = []
    used_subdocs = set()

    for flag in flags:
        base = flag.replace("ENABLE_", "")
        placeholder = f"{base}_SUBDOC"
        if placeholder in subdocs:
            sid = base.lower()
            # 如果有中文映射，使用中文标题
            title = SECTION_ID_TO_CHINESE.get(sid, base.replace("_", " ").title())
            sections.append({
                "id": sid,
                "placeholder": placeholder,
                "flag_name": flag,
                "title": title,
            })
            used_subdocs.add(placeholder)

    # 处理未配对的 subdocs
    for sd in subdocs:
        if sd not in used_subdocs and sd != "APPENDICES_SUBDOC":
            base = sd.replace("_SUBDOC", "")
            sid = base.lower()
            title = SECTION_ID_TO_CHINESE.get(sid, base.replace("_", " ").title())
            sections.append({
                "id": sid,
                "placeholder": sd,
                "flag_name": f"ENABLE_{base.upper()}",
                "title": title,
            })

    return sections


def auto_generate_from_template(
    template_path: str,
    context: dict,
    dry_run: bool = False,
) -> dict:
    """根据模板 prompts 自动生成内容，返回简化内容描述。"""
    # 解析模板完整结构
    script_dir = Path(__file__).resolve().parent
    sys.path.insert(0, str(script_dir))
    try:
        from analyze_template import analyze_template
        template_structure = analyze_template(template_path)
    except ImportError:
        template_structure = {
            "prompts": extract_prompts_from_template(template_path),
            "subdoc_placeholders": [],
            "conditional_flags": [],
        }

    prompts = template_structure.get("prompts", [])
    subdocs = template_structure.get("subdoc_placeholders", [])
    flags = template_structure.get("conditional_flags", [])

    if not prompts:
        raise ValueError(f"模板 {template_path} 中没有发现 PROMPT 注释，无法自动生成")

    # 从 flags + subdocs 构建 sections
    template_sections = _build_sections_from_flags(flags, subdocs)

    # 建立 prompt target -> template section 的映射
    section_meta_by_prompt: dict = {}
    for p in prompts:
        if p["level"] != "section":
            continue
        target = p["target"]
        matched = None

        # 精确匹配：target == title 或 target == id
        for ts in template_sections:
            if target.lower() == ts.get("title", "").lower() or target.lower() == ts.get("id", "").lower():
                matched = ts
                break

        # 模糊匹配
        if not matched:
            for ts in template_sections:
                title = ts.get("title", "").lower()
                tid = ts.get("id", "").lower()
                tplaceholder = ts.get("placeholder", "").lower()
                if (target.lower() in title or title in target.lower() or
                    target.lower() in tid or tid in target.lower() or
                    target.lower() in tplaceholder or tplaceholder in target.lower()):
                    matched = ts
                    break

        section_meta_by_prompt[target] = {
            "prompt": p,
            "template_section": matched,
        }

    if not section_meta_by_prompt:
        raise ValueError("模板中没有 section-level 的 PROMPT，无法自动生成")

    # 构建文档结构
    document_structure = []
    for target, meta in section_meta_by_prompt.items():
        ts = meta["template_section"]
        doc_id = ts["id"] if ts else target
        doc_title = ts["title"] if ts else target
        document_structure.append({"id": doc_id, "title": doc_title})

    # 导入 LLM 客户端和 Markdown 转换器
    from llm_client import generate_section_content
    from markdown_to_blocks import markdown_to_blocks

    sections = []
    for idx, (target, meta) in enumerate(section_meta_by_prompt.items()):
        prompt_meta = meta["prompt"]
        ts = meta["template_section"]
        prompt_text = prompt_meta["prompt"]
        mode = prompt_meta["mode"]

        display_name = ts["title"] if ts else target
        print(f"  生成章节: {display_name} (mode={mode})")

        # section 之间添加延迟，避免并发请求上游限流
        if idx > 0 and not dry_run:
            import time
            time.sleep(1)

        if dry_run:
            blocks = [{"type": "paragraph", "text": f"[DRY-RUN] {prompt_text}"}]
        else:
            try:
                markdown = generate_section_content(
                    prompt=prompt_text,
                    target=display_name,
                    context=context,
                    document_structure=document_structure,
                )
                blocks = markdown_to_blocks(markdown)
                if not blocks:
                    blocks = [{"type": "paragraph", "text": markdown}]
            except Exception as e:
                print(f"    警告: 生成失败 ({e})，使用空内容")
                blocks = []

        section_data = {
            "name": display_name,
            "blocks": blocks,
            "enabled": True,
        }
        # 如果匹配到了模板 section，使用其 id/placeholder/flag_name
        if ts:
            section_data["id"] = ts.get("id")
            section_data["placeholder"] = ts.get("placeholder")
            section_data["flag_name"] = ts.get("flag_name")

        sections.append(section_data)

    return {
        "title": context.get("PROJECT_NAME", "未命名报告"),
        "org": context.get("APPLICANT_ORG", ""),
        "leader": context.get("PROJECT_LEADER", ""),
        "period": context.get("PROJECT_PERIOD", ""),
        "extra_context": context,
        "sections": sections,
        "attachments": [],
    }


# ── 主函数 ─────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="从简化内容描述构建 payload")
    parser.add_argument("--content", help="内容描述 JSON 文件路径")
    parser.add_argument("--output", help="输出 payload 路径")
    parser.add_argument("--template", help="模板文件路径（可选，用于交叉校验和 prompt 提取）")
    parser.add_argument("--list-prompts", action="store_true", help="仅列出模板中的 PROMPT 注释")
    parser.add_argument("--auto-generate", action="store_true", help="根据模板 prompts 自动调用 LLM 生成内容")
    parser.add_argument("--context", help="JSON 格式的上下文变量（用于 --auto-generate）")
    parser.add_argument("--dry-run", action="store_true", help="模拟生成，不实际调用 LLM")
    args = parser.parse_args()

    # 模式 1：仅列出 prompts
    if args.list_prompts:
        if not args.template:
            print("错误: --list-prompts 需要 --template")
            sys.exit(1)
        list_prompts(args.template)
        return

    # 模式 2：自动生成
    if args.auto_generate:
        if not args.template:
            print("错误: --auto-generate 需要 --template")
            sys.exit(1)
        if not args.output:
            print("错误: --auto-generate 需要 --output")
            sys.exit(1)

        context = {}
        if args.context:
            try:
                context = json.loads(args.context)
            except json.JSONDecodeError as e:
                print(f"错误: --context JSON 解析失败: {e}")
                sys.exit(1)

        print(f"从模板自动生成内容: {args.template}")
        content = auto_generate_from_template(args.template, context, dry_run=args.dry_run)

        try:
            payload = build_payload(content, template_path=args.template)
        except ValueError as e:
            print(f"错误: {e}", file=sys.stderr)
            sys.exit(1)

        Path(args.output).parent.mkdir(parents=True, exist_ok=True)
        with open(args.output, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)

        print(f"\nPayload 已生成: {args.output}")
        print(f"  context: {len(payload['context'])} 个字段")
        print(f"  sections: {len(payload['sections'])} 个")
        print(f"  attachments: {len(payload['attachments'])} 个")

        # 可选：用 report-engine 校验
        try:
            import subprocess
            result = subprocess.run(
                ["report-engine", "check-template",
                 "--template", args.template,
                 "--payload", args.output],
                capture_output=True, text=True
            )
            print(result.stdout)
            if result.returncode != 0:
                print(f"校验失败: {result.stderr}", file=sys.stderr)
        except FileNotFoundError:
            pass  # report-engine CLI 未安装，跳过校验
        return

    # 模式 3：标准模式（从 content.json 构建）
    if not args.content:
        print("错误: 需要 --content 或 --auto-generate")
        sys.exit(1)
    if not args.output:
        print("错误: 需要 --output")
        sys.exit(1)

    content = json.loads(Path(args.content).read_text(encoding="utf-8"))

    try:
        payload = build_payload(content, template_path=args.template)
    except ValueError as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)

    Path(args.output).parent.mkdir(parents=True, exist_ok=True)
    with open(args.output, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    print(f"Payload 已生成: {args.output}")
    print(f"  context: {len(payload['context'])} 个字段")
    print(f"  sections: {len(payload['sections'])} 个")
    print(f"  attachments: {len(payload['attachments'])} 个")

    # 可选：用 report-engine 校验
    if args.template:
        try:
            import subprocess
            result = subprocess.run(
                ["report-engine", "check-template",
                 "--template", args.template,
                 "--payload", args.output],
                capture_output=True, text=True
            )
            print(result.stdout)
            if result.returncode != 0:
                print(f"校验失败: {result.stderr}", file=sys.stderr)
        except FileNotFoundError:
            pass


if __name__ == "__main__":
    main()
